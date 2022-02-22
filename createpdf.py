# Creating a demp invoice structue using python libraries. 

# import aspose.words as aw 

# # create document object 
# doc = aw.Document()

# # create a document buiilder class
# builder = aw.DocumentBuilder(doc)

# #add text to the document
# builder.write('Hello new docx')

# #save doc
# doc.save('InvoiceFiles\demo.docx')

""" Above shown library is not for free. so moving to the python-docx"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
for i in range(5):
    document.add_paragraph(" ")

 # adding invoice aligned at center.   
invoice_para = document.add_paragraph("")
run = invoice_para.add_run("INVOICE")
invoice_para.alignment = 1
run.bold = True
# # or use the following ::::
run.font.size = Pt(25)
run.font.name = 'Times New Roman'




document.save('InvoiceFiles\demo2.docx')

